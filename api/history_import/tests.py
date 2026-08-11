from datetime import date, timedelta
from io import BytesIO
from unittest.mock import patch

import pytest
from django.core.files.uploadedfile import SimpleUploadedFile
from rest_framework.test import APIClient

from accounts.models import User
from cars.models import Car
from expenses.models import Expense
from history_import.duplicates import find_duplicate
from history_import.extraction import ExtractionError, _normalize_record, extract_records
from services.models import ServiceRecord


@pytest.fixture
def owner(db):
    return User.objects.create_user(email="owner@example.com", password="str0ng-pass-123")


@pytest.fixture
def car(owner):
    return Car.objects.create(owner=owner, make="Toyota", model="Corolla", current_odometer_km=40000)


@pytest.fixture
def client(owner):
    api = APIClient()
    api.force_authenticate(owner)
    return api


class TestNormalizeRecord:
    """See #103."""

    def test_known_category_hint_maps_to_a_real_service_type(self):
        record = _normalize_record({"kind": "service", "date": "2025-01-01", "category_hint": "oil_change", "description": "Oil + filter"})
        assert record["service_type"] == "oil_change"
        assert record["description"] == "Oil + filter"

    def test_unmatched_category_hint_folds_into_the_description(self):
        # No ServiceRecord.service_type choice for "suspension" -- must not
        # be silently dropped.
        record = _normalize_record({"kind": "service", "date": "2025-01-01", "category_hint": "suspension", "description": "Replaced bushings"})
        assert record["service_type"] == "other"
        assert record["description"] == "Suspension service — Replaced bushings"

    def test_part_purchase_has_no_service_type(self):
        record = _normalize_record({"kind": "part_purchase", "date": "2025-01-01", "category_hint": "engine", "description": "Turbo"})
        assert record["service_type"] is None
        assert record["kind"] == "part_purchase"

    def test_missing_kind_defaults_to_service(self):
        record = _normalize_record({"date": "2025-01-01", "description": "Something"})
        assert record["kind"] == "service"

    def test_non_numeric_cost_becomes_none_rather_than_erroring(self):
        record = _normalize_record({"kind": "service", "date": "2025-01-01", "cost": "not a number"})
        assert record["cost"] is None


class TestExtractRecords:
    """See #103."""

    def test_unsupported_file_type_raises(self):
        upload = SimpleUploadedFile("history.txt", b"whatever", content_type="text/plain")
        with pytest.raises(ExtractionError):
            extract_records(upload)

    @patch("history_import.extraction._call_gemini", return_value=[{"kind": "service", "date": "2025-01-01", "description": "Oil change"}])
    def test_xlsx_extracts_text_before_calling_gemini(self, mock_call):
        import openpyxl

        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.append(["Date", "Description", "Cost"])
        sheet.append(["2025-01-01", "Oil change", 45])
        buffer = BytesIO()
        workbook.save(buffer)
        buffer.seek(0)
        upload = SimpleUploadedFile("history.xlsx", buffer.read(), content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        records = extract_records(upload)

        assert len(records) == 1
        # The text passed to Gemini should include the sheet's actual data.
        contents_arg = mock_call.call_args[0][0]
        assert "Oil change" in contents_arg[0]

    @patch("history_import.extraction._call_gemini", return_value=[{"kind": "service", "date": "2025-01-01", "description": "Oil change"}])
    def test_docx_extracts_text_before_calling_gemini(self, mock_call):
        import docx

        document = docx.Document()
        document.add_paragraph("Garage: QuickLube")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "2025-01-01"
        table.rows[0].cells[1].text = "Oil change — $45"
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        upload = SimpleUploadedFile("history.docx", buffer.read(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        records = extract_records(upload)

        assert len(records) == 1
        contents_arg = mock_call.call_args[0][0]
        assert "QuickLube" in contents_arg[0]
        assert "Oil change" in contents_arg[0]

    def test_docx_with_no_text_raises(self):
        import docx

        document = docx.Document()
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        upload = SimpleUploadedFile("empty.docx", buffer.read(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        with pytest.raises(ExtractionError):
            extract_records(upload)

    def test_xlsx_with_no_data_raises(self):
        import openpyxl

        workbook = openpyxl.Workbook()
        buffer = BytesIO()
        workbook.save(buffer)
        buffer.seek(0)
        upload = SimpleUploadedFile("empty.xlsx", buffer.read(), content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        with pytest.raises(ExtractionError):
            extract_records(upload)


@pytest.mark.django_db
class TestFindDuplicate:
    """See #103 — e.g. an oil change already logged for 2025 shouldn't get
    re-added because a newly-imported document also mentions one from
    around then, even without an exact date/cost match."""

    def test_same_service_type_within_the_window_is_a_duplicate(self, car):
        ServiceRecord.objects.create(car=car, service_type="oil_change", service_date=date(2025, 3, 1), odometer_km=1000)

        record = {"kind": "service", "service_type": "oil_change", "date": date(2025, 3, 20), "cost": None, "vendor": ""}

        assert find_duplicate(car, record) is not None

    def test_different_service_type_is_not_a_duplicate(self, car):
        ServiceRecord.objects.create(car=car, service_type="oil_change", service_date=date(2025, 3, 1), odometer_km=1000)

        record = {"kind": "service", "service_type": "brakes", "date": date(2025, 3, 20), "cost": None, "vendor": ""}

        assert find_duplicate(car, record) is None

    def test_same_service_type_outside_the_window_is_not_a_duplicate(self, car):
        ServiceRecord.objects.create(car=car, service_type="oil_change", service_date=date(2025, 3, 1), odometer_km=1000)

        far_date = date(2025, 3, 1) + timedelta(days=200)
        record = {"kind": "service", "service_type": "oil_change", "date": far_date, "cost": None, "vendor": ""}

        assert find_duplicate(car, record) is None

    def test_part_purchase_matches_on_cost_within_the_window(self, car):
        Expense.objects.create(car=car, category="modification_parts", amount=120, expense_date=date(2025, 5, 1))

        record = {"kind": "part_purchase", "date": date(2025, 5, 10), "cost": 120, "vendor": ""}

        assert find_duplicate(car, record) is not None

    def test_part_purchase_matches_on_vendor_within_the_window(self, car):
        Expense.objects.create(car=car, category="modification_parts", amount=999, expense_date=date(2025, 5, 1), vendor="AutoParts Co")

        record = {"kind": "part_purchase", "date": date(2025, 5, 10), "cost": None, "vendor": "AutoParts Co"}

        assert find_duplicate(car, record) is not None

    def test_part_purchase_with_neither_cost_nor_vendor_matching_is_not_a_duplicate(self, car):
        Expense.objects.create(car=car, category="modification_parts", amount=120, expense_date=date(2025, 5, 1), vendor="AutoParts Co")

        record = {"kind": "part_purchase", "date": date(2025, 5, 10), "cost": 50, "vendor": "Different Vendor"}

        assert find_duplicate(car, record) is None

    def test_unparseable_date_is_not_a_duplicate_rather_than_erroring(self, car):
        record = {"kind": "service", "service_type": "oil_change", "date": "not a date", "cost": None, "vendor": ""}

        assert find_duplicate(car, record) is None


@pytest.mark.django_db
class TestServiceHistoryExtractView:
    """See #103."""

    @patch("history_import.views.extract_records", return_value=[{"kind": "service", "date": "2025-01-01", "description": "Oil change", "vendor": "", "cost": None, "odometer_km": None, "service_type": "oil_change"}])
    def test_returns_extracted_records(self, mock_extract, client, car):
        upload = SimpleUploadedFile("history.pdf", b"%PDF-fake", content_type="application/pdf")

        response = client.post("/api/v1/history-import/extract/", {"car": str(car.id), "file": upload}, format="multipart")

        assert response.status_code == 200
        assert response.data["records"][0]["description"] == "Oil change"
        assert response.data["records"][0]["possible_duplicate"] is False

    @patch("history_import.views.extract_records", return_value=[{"kind": "service", "date": "2025-01-01", "description": "Oil change", "vendor": "", "cost": None, "odometer_km": None, "service_type": "oil_change"}])
    def test_flags_a_likely_duplicate(self, mock_extract, client, car):
        ServiceRecord.objects.create(car=car, service_type="oil_change", service_date=date(2025, 1, 5), odometer_km=1000)
        upload = SimpleUploadedFile("history.pdf", b"%PDF-fake", content_type="application/pdf")

        response = client.post("/api/v1/history-import/extract/", {"car": str(car.id), "file": upload}, format="multipart")

        assert response.data["records"][0]["possible_duplicate"] is True

    def test_rejects_a_car_the_owner_doesnt_own(self, client):
        other_owner = User.objects.create_user(email="other@example.com", password="str0ng-pass-123")
        other_car = Car.objects.create(owner=other_owner, make="Honda", model="Civic")
        upload = SimpleUploadedFile("history.pdf", b"%PDF-fake", content_type="application/pdf")

        response = client.post("/api/v1/history-import/extract/", {"car": str(other_car.id), "file": upload}, format="multipart")

        assert response.status_code == 404

    def test_rejects_an_unsupported_file_type(self, client, car):
        upload = SimpleUploadedFile("history.txt", b"plain text", content_type="text/plain")

        response = client.post("/api/v1/history-import/extract/", {"car": str(car.id), "file": upload}, format="multipart")

        assert response.status_code == 400

    def test_requires_authentication(self, car):
        api = APIClient()
        upload = SimpleUploadedFile("history.pdf", b"%PDF-fake", content_type="application/pdf")

        response = api.post("/api/v1/history-import/extract/", {"car": str(car.id), "file": upload}, format="multipart")

        assert response.status_code == 401


@pytest.mark.django_db
class TestServiceHistoryImportView:
    """See #103."""

    def test_creates_a_service_record(self, client, car):
        payload = {
            "car": str(car.id),
            "records": [{"kind": "service", "date": "2025-01-01", "vendor": "QuickLube", "description": "Oil change", "cost": "45.00", "service_type": "oil_change"}],
        }

        response = client.post("/api/v1/history-import/confirm/", payload, format="json")

        assert response.status_code == 201
        assert response.data["service_records_created"] == 1
        record = ServiceRecord.objects.get(car=car)
        assert record.garage_name == "QuickLube"
        assert record.cost == 45

    def test_creates_an_expense_for_a_part_purchase(self, client, car):
        payload = {
            "car": str(car.id),
            "records": [{"kind": "part_purchase", "date": "2025-01-01", "vendor": "AutoParts Co", "description": "Turbo", "cost": "500.00"}],
        }

        response = client.post("/api/v1/history-import/confirm/", payload, format="json")

        assert response.status_code == 201
        assert response.data["expenses_created"] == 1
        expense = Expense.objects.get(car=car)
        assert expense.category == "modification_parts"
        assert expense.amount == 500

    def test_skips_a_duplicate_as_a_hard_backstop(self, client, car):
        ServiceRecord.objects.create(car=car, service_type="oil_change", service_date=date(2025, 1, 5), odometer_km=1000)
        payload = {
            "car": str(car.id),
            "records": [{"kind": "service", "date": "2025-01-10", "vendor": "", "description": "Oil change", "cost": None, "service_type": "oil_change"}],
        }

        response = client.post("/api/v1/history-import/confirm/", payload, format="json")

        assert response.status_code == 201
        assert response.data["service_records_created"] == 0
        assert response.data["duplicates_skipped"] == 1
        assert ServiceRecord.objects.filter(car=car).count() == 1

    def test_missing_odometer_defaults_to_zero_without_moving_the_cars_reading_backwards(self, client, car):
        payload = {
            "car": str(car.id),
            "records": [{"kind": "service", "date": "2025-01-01", "vendor": "", "description": "Old service", "cost": None, "service_type": "minor_service"}],
        }

        client.post("/api/v1/history-import/confirm/", payload, format="json")

        car.refresh_from_db()
        assert car.current_odometer_km == 40000

    def test_rejects_a_car_the_owner_doesnt_own(self, client):
        other_owner = User.objects.create_user(email="other@example.com", password="str0ng-pass-123")
        other_car = Car.objects.create(owner=other_owner, make="Honda", model="Civic")
        payload = {"car": str(other_car.id), "records": [{"kind": "service", "date": "2025-01-01", "description": "x"}]}

        response = client.post("/api/v1/history-import/confirm/", payload, format="json")

        assert response.status_code == 404

    def test_requires_authentication(self, car):
        api = APIClient()
        payload = {"car": str(car.id), "records": [{"kind": "service", "date": "2025-01-01", "description": "x"}]}

        response = api.post("/api/v1/history-import/confirm/", payload, format="json")

        assert response.status_code == 401
