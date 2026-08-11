"""
Turns an uploaded document (PDF, .docx, .xlsx) into proposed ServiceRecord/
Expense rows — see #103. Extraction only proposes; nothing is saved until
the owner reviews and confirms via ServiceHistoryImportView.

PDF goes to Gemini as-is (native multimodal document understanding, no
separate parsing needed). .docx/.xlsx have no such native support, so their
text is extracted first (python-docx / openpyxl) and sent as plain text —
same downstream prompt/schema either way. The uploaded file itself is only
ever read into memory here, never written to disk (see #33's no-local-disk-
uploads convention) — nothing about it needs to persist past extraction.
"""
import json

from django.conf import settings

from utils import Constants

# service_type choices Gemini can map straight to a ServiceRecord field.
# Anything else it proposes (e.g. "suspension", "engine" — no dedicated
# choice exists for either) falls back to "other" with the category folded
# into the description instead of silently dropped.
_SERVICE_TYPE_HINTS = {
    "oil_change": Constants.SERVICE_TYPE_OIL_CHANGE,
    "brakes": Constants.SERVICE_TYPE_BRAKES,
    "tyres": Constants.SERVICE_TYPE_TYRES,
    "battery": Constants.SERVICE_TYPE_BATTERY,
    "minor_service": Constants.SERVICE_TYPE_MINOR,
    "major_service": Constants.SERVICE_TYPE_MAJOR,
}

RESPONSE_SCHEMA = {
    "type": "object",
    "properties": {
        "records": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "kind": {"type": "string", "enum": ["service", "part_purchase"]},
                    "date": {"type": "string", "description": "YYYY-MM-DD; best guess if the document is ambiguous"},
                    "vendor": {"type": "string", "description": "Garage name or parts vendor"},
                    "description": {"type": "string", "description": "What work was done or what was purchased"},
                    "cost": {"type": "number"},
                    "odometer_km": {"type": "integer", "description": "Only if actually stated in the document"},
                    "category_hint": {
                        "type": "string",
                        "description": (
                            "One of: oil_change, brakes, tyres, battery, minor_service, "
                            "major_service, suspension, engine, other"
                        ),
                    },
                },
                "required": ["kind", "date", "description"],
            },
        },
    },
    "required": ["records"],
}

EXTRACTION_PROMPT = (
    "This document is a car's service/maintenance history — an invoice, receipt, "
    "or log the owner kept before using this app. Extract every distinct service "
    "or part purchase as its own record. For each: the date (best guess if "
    "unclear), garage/vendor name, a short description of the work done or item "
    "bought, the cost if stated, the odometer reading only if it's actually "
    "written down, and a category_hint. Use 'part_purchase' as kind only for "
    "buying/sourcing a part, not labor. If a record doesn't fit the maintenance "
    "categories, use category_hint 'other' with the specifics in the description "
    "rather than guessing a mismatched category. If the document has nothing "
    "resembling car service/maintenance/parts records, return an empty list."
)


class ExtractionError(Exception):
    """Raised when the document can't be turned into records at all —
    unsupported/corrupt file, or the model returned something unusable."""


def extract_records(uploaded_file):
    """
    :param uploaded_file: a Django UploadedFile (PDF/.docx/.xlsx).
    :returns: list of dicts — kind, date, vendor, description, cost,
        odometer_km, service_type (None for a part_purchase row).
    :raises ExtractionError: unsupported type, unreadable file, or the
        model call/response failed.
    """
    name = (uploaded_file.name or "").lower()
    if name.endswith(".pdf"):
        contents = [_gemini_pdf_part(uploaded_file)]
    elif name.endswith(".docx"):
        contents = [_extract_docx_text(uploaded_file)]
    elif name.endswith(".xlsx"):
        contents = [_extract_xlsx_text(uploaded_file)]
    else:
        raise ExtractionError("Unsupported file type — upload a PDF, .docx, or .xlsx file.")

    records = _call_gemini(contents)
    return [_normalize_record(r) for r in records if isinstance(r, dict)]


def _gemini_pdf_part(uploaded_file):
    from google.genai import types

    return types.Part.from_bytes(data=uploaded_file.read(), mime_type="application/pdf")


def _extract_docx_text(uploaded_file):
    import docx

    try:
        document = docx.Document(uploaded_file)
    except Exception as err:
        raise ExtractionError("Couldn't read that Word document.") from err

    lines = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text for cell in row.cells))

    text = "\n".join(lines)
    if not text.strip():
        raise ExtractionError("Couldn't find any text in that document.")
    return text


def _extract_xlsx_text(uploaded_file):
    import openpyxl

    try:
        workbook = openpyxl.load_workbook(uploaded_file, data_only=True, read_only=True)
    except Exception as err:
        raise ExtractionError("Couldn't read that spreadsheet.") from err

    lines = []
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                lines.append(", ".join("" if cell is None else str(cell) for cell in row))

    text = "\n".join(lines)
    if not text.strip():
        raise ExtractionError("Couldn't find any data in that spreadsheet.")
    return text


def _call_gemini(contents):
    from google import genai
    from google.genai import types

    if not settings.GEMINI_API_KEY:
        raise ExtractionError("Document import isn't configured on this server.")

    client = genai.Client(api_key=settings.GEMINI_API_KEY)
    config = types.GenerateContentConfig(
        system_instruction=EXTRACTION_PROMPT,
        response_mime_type="application/json",
        response_schema=RESPONSE_SCHEMA,
        temperature=0.1,
    )
    try:
        response = client.models.generate_content(
            model=settings.GEMINI_MODEL, contents=contents, config=config,
        )
        data = json.loads(response.text)
    except ExtractionError:
        raise
    except Exception as err:
        raise ExtractionError("Couldn't read that document. Please try again.") from err

    records = data.get("records") if isinstance(data, dict) else None
    if not isinstance(records, list):
        raise ExtractionError("Couldn't read that document. Please try again.")
    return records


def _normalize_record(raw):
    """Maps one raw Gemini record into the shape the frontend review step
    and confirm endpoint both expect — resolves category_hint down to a
    real ServiceRecord.service_type, folding it into the description
    instead when there's no matching choice (e.g. suspension/engine)."""
    kind = raw.get("kind") if raw.get("kind") in ("service", "part_purchase") else "service"
    category_hint = (raw.get("category_hint") or "").strip().lower()
    description = (raw.get("description") or "").strip()

    service_type = None
    if kind == "service":
        service_type = _SERVICE_TYPE_HINTS.get(category_hint, Constants.SERVICE_TYPE_OTHER)
        if category_hint and category_hint not in _SERVICE_TYPE_HINTS and category_hint != "other":
            label = category_hint.replace("_", " ").capitalize()
            description = f"{label} service" + (f" — {description}" if description else "")

    cost = raw.get("cost")
    odometer_km = raw.get("odometer_km")
    return {
        "kind": kind,
        "date": raw.get("date") or "",
        "vendor": (raw.get("vendor") or "").strip(),
        "description": description,
        "cost": cost if isinstance(cost, (int, float)) else None,
        "odometer_km": odometer_km if isinstance(odometer_km, (int, float)) else None,
        "service_type": service_type,
    }
